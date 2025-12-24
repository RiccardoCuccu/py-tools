"""
Text Extraction Module - STEP 1

Handles extraction of text from various document formats (.docx, .pdf, .txt)
with support for partial page extraction.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    import fitz
except ImportError as e:
    print(f"Error: Missing required library - {e}")
    print("pip install python-docx PyMuPDF")
    sys.exit(1)


class TextExtractor:
    def __init__(self, doc_path, extract_pages=None, page_position='middle'):
        """Initialize text extractor with document path and extraction options"""
        self.doc_path = Path(doc_path)
        self.extract_pages = extract_pages
        self.page_position = page_position
    
    def extract_text(self):
        """Extract text based on file extension (dispatcher method)"""
        if self.doc_path.suffix.lower() == '.docx':
            return self._extract_from_docx()
        elif self.doc_path.suffix.lower() == '.pdf':
            return self._extract_from_pdf()
        elif self.doc_path.suffix.lower() == '.txt':
            text = self._extract_from_txt()
            if text:
                return text
            else:
                print(f"✗ Error reading text file")
                sys.exit(1)
        else:
            print(f"✗ Error: Unsupported file format: {self.doc_path.suffix}")
            print("Supported formats: .docx, .pdf, .txt")
            sys.exit(1)
    
    def _extract_page_range(self, text, chars_per_page=3000):
        """Extract specific page range from text based on config"""
        if self.extract_pages is None:
            return text
        
        total_chars = len(text)
        chars_to_extract = self.extract_pages * chars_per_page
        
        if self.page_position == 'start':
            start_char = 0
        elif self.page_position == 'end':
            start_char = max(0, total_chars - chars_to_extract)
        else:  # middle
            start_char = max(0, (total_chars - chars_to_extract) // 2)
        
        end_char = min(start_char + chars_to_extract, total_chars)
        return text[start_char:end_char]
    
    def _extract_from_docx(self):
        """Extract text from a Word document (optionally from specific pages)"""
        print(f"\n[1/5] Extracting text from document: {self.doc_path.name}")
        try:
            doc = Document(str(self.doc_path))
            all_paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = '\n'.join(all_paragraphs)
            
            if self.extract_pages is None:
                print(f"✓ Extracted {len(full_text)} characters from entire document")
                return full_text
            else:
                # Extract specific pages
                CHARS_PER_PAGE = 3000
                total_chars = len(full_text)
                total_pages = max(1, total_chars // CHARS_PER_PAGE)
                
                text = self._extract_page_range(full_text, CHARS_PER_PAGE)
                print(f"✓ Extracted {len(text)} characters from {self.extract_pages} pages ({self.page_position}) out of ~{total_pages} total pages")
                return text
        except Exception as e:
            print(f"✗ Error reading document: {e}")
            sys.exit(1)
    
    def _extract_from_pdf(self):
        """Extract text from a PDF document"""
        print(f"\n[1/5] Extracting text from PDF: {self.doc_path.name}")
        try:
            with fitz.open(str(self.doc_path)) as doc:  # type: ignore[attr-defined]
                full_text = "".join(page.get_text() for page in doc)
                num_pages = len(doc)
            
            if self.extract_pages is None:
                print(f"✓ Extracted {len(full_text)} characters from {num_pages} pages")
                return full_text
            else:
                text = self._extract_page_range(full_text)
                print(f"✓ Extracted {len(text)} characters from {self.extract_pages} pages ({self.page_position}) out of {num_pages} total pages")
                return text
        except Exception as e:
            print(f"✗ Error reading PDF: {e}")
            sys.exit(1)
    
    def _extract_from_txt(self, file_path=None):
        """Extract text from a plain text file (optionally from specific pages)"""
        if file_path is None:
            file_path = self.doc_path
            
        print(f"\n[1/5] Extracting text from text file: {Path(file_path).name}")
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
            
            # Apply page extraction only for main document
            if file_path == self.doc_path:
                if self.extract_pages is not None:
                    CHARS_PER_PAGE = 3000
                    total_chars = len(full_text)
                    total_pages = max(1, total_chars // CHARS_PER_PAGE)
                    
                    text = self._extract_page_range(full_text, CHARS_PER_PAGE)
                    print(f"✓ Extracted {len(text)} characters from {self.extract_pages} pages ({self.page_position}) out of ~{total_pages} total pages")
                    return text
                else:
                    print(f"✓ Extracted {len(full_text)} characters from entire file")
                    return full_text
            else:
                # For reference files, always return full text
                return full_text
        except Exception as e:
            print(f"✗ Error reading text file {file_path}: {e}")
            return None
    
    def extract_from_file(self, file_path):
        """Extract text from any supported file format (used for reference files)"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == '.docx':
            try:
                doc = Document(str(file_path))
                return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                print(f"  Warning: Failed to read {file_path.name}: {e}")
                return None
        elif suffix == '.pdf':
            try:
                with fitz.open(str(file_path)) as pdf_doc:  # type: ignore[attr-defined]
                    return "".join(page.get_text() for page in pdf_doc)
            except Exception as e:
                print(f"  Warning: Failed to read {file_path.name}: {e}")
                return None
        elif suffix == '.txt':
            return self._extract_from_txt(file_path)
        else:
            print(f"  Warning: Unsupported format {suffix} for {file_path.name}")
            return None